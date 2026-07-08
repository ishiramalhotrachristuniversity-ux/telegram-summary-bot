from http.server import BaseHTTPRequestHandler
import os
import json
import requests
import google.generativeai as genai
from io import BytesIO
from docx import Document

# تنظیمات کلاینت جمینای
genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
TELEGRAM_URL = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"

# فراخوانی آیدی گروه و حذف فاصله‌های احتمالی
GROUP_CHAT_ID = str(os.environ.get("GROUP_CHAT_ID", "")).strip()

def get_best_model():
    models = genai.list_models()
    for m in models:
        if 'generateContent' in m.supported_generation_methods and 'gemini' in m.name:
            return genai.GenerativeModel(m.name)
    return genai.GenerativeModel('gemini-pro')

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers.get('Content-Length', 0))
        post_data = self.rfile.read(content_length)
        update = json.loads(post_data)

        if 'message' in update and 'document' in update['message']:
            chat_id = update['message']['chat']['id']
            document = update['message']['document']
            
            if document.get('file_name', '').endswith('.docx'):
                try:
                    file_id = document['file_id']
                    file_info = requests.get(f"{TELEGRAM_URL}/getFile?file_id={file_id}").json()
                    file_path = file_info['result']['file_path']
                    download_url = f"https://api.telegram.org/file/bot{TELEGRAM_TOKEN}/{file_path}"
                    file_content = requests.get(download_url).content
                    
                    doc = Document(BytesIO(file_content))
                    text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    model = get_best_model()
                    
                    prompt = (
                        "گزارش زیر را در قالب یک فایل ورد با فرمت دقیق زیر خلاصه کن.\n\n"
                        "قوانین اجباری:\n"
                        "1. نام رویداد را دقیقاً از متن استخراج کن.\n"
                        "2. اگر تاریخ رویداد در متن وجود داشت، همان را دقیقاً استخراج کن و روبروی بخش تاریخ رویداد بنویس.\n"
                        "3. کل بخش «خلاصه رویداد» باید بین 5 تا 7 خط باشد.\n"
                        "4. لحن روایی گزارش حفظ شود.\n"
                        "5. چیزی از خودت اضافه نکن و فقط از اطلاعات موجود در متن استفاده کن.\n\n"
                        "فرمت خروجی (دقیقاً به همین شکل شروع شود):\n"
                        "نام رویداد: [عنوان دقیق رویداد]\n"
                        "تاریخ رویداد: [تاریخ رویداد یا عبارت ذکر نشده]\n"
                        "خلاصه رویداد:\n"
                        "[خلاصه‌ای توصیفی و روایی در 5 تا 7 خط]\n\n"
                        "ارزش محوری: [یک جمله کوتاه]\n\n"
                        f"متن گزارش:\n{text}"
                    )
                    
                    response = model.generate_content(prompt)
                    
                    # --- استخراج هوشمند نام رویداد برای اسم فایل ---
                    event_name = "گزارش_رویداد" # نام پیش‌فرض در صورت پیدا نشدن
                    for line in response.text.split('\n'):
                        if line.strip().startswith("نام رویداد:"):
                            extracted_name = line.split("نام رویداد:")[1].strip()
                            # حذف براکت‌ها اگر هوش مصنوعی اشتباهاً گذاشته بود
                            extracted_name = extracted_name.replace("[", "").replace("]", "").strip()
                            if extracted_name:
                                event_name = extracted_name
                            break
                            
                    # پاک‌سازی کاراکترهای غیرمجاز ویندوز/تلگرام برای اسم فایل
                    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
                    for char in invalid_chars:
                        event_name = event_name.replace(char, ' ')
                    
                    event_name = event_name.strip()
                    if not event_name:
                        event_name = "گزارش_رویداد"
                        
                    final_file_name = f"{event_name}.docx"
                    # -----------------------------------------------
                    
                    new_doc = Document()
                    new_doc.add_paragraph(response.text)
                    
                    output = BytesIO()
                    new_doc.save(output)
                    output.seek(0)
                    
                    if GROUP_CHAT_ID:
                        # استفاده از نام استخراج شده برای فایل ارسالی
                        files = {'document': (final_file_name, output, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                        send_response = requests.post(f"{TELEGRAM_URL}/sendDocument", data={'chat_id': GROUP_CHAT_ID}, files=files)
                        
                        if send_response.status_code == 200:
                            requests.post(f"{TELEGRAM_URL}/sendMessage", json={"chat_id": chat_id, "text": f"✅ گزارش با موفقیت به گروه ارسال شد.\nنام فایل ثبت شده: {final_file_name}"})
                        else:
                            error_desc = send_response.json().get('description', 'دلیل نامشخص')
                            requests.post(f"{TELEGRAM_URL}/sendMessage", json={"chat_id": chat_id, "text": f"⚠️ خطا در ارسال به گروه: {error_desc}"})
                    else:
                        requests.post(f"{TELEGRAM_URL}/sendMessage", json={"chat_id": chat_id, "text": "⚠️ متغیر GROUP_CHAT_ID در تنظیمات ورسل خالی است!"})
                        
                except Exception as e:
                    requests.post(f"{TELEGRAM_URL}/sendMessage", json={"chat_id": chat_id, "text": f"خطا در پردازش: {str(e)}"})
            
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b"OK")
